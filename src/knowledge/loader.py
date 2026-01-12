"""
Knowledge base document loader for CyraSoul.
Supports multiple document formats: Markdown, PDF, Word.
"""

import os
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Generator
from dataclasses import dataclass, field
import tiktoken

# PDF support
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# Word document support
try:
    from docx import Document as WordDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@dataclass
class DocumentChunk:
    """Represents a chunk of a document."""
    id: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class LoadedDocument:
    """Represents a fully loaded document."""
    id: str
    filename: str
    content: str
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class KnowledgeLoader:
    """
    Loader for knowledge base documents.
    Handles various file formats and chunking.
    """

    SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        min_chunk_size: int = 100
    ):
        """
        Initialize the knowledge loader.

        Args:
            chunk_size: Target chunk size in tokens.
            chunk_overlap: Overlap between chunks in tokens.
            min_chunk_size: Minimum chunk size in tokens.
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size

        # Initialize tokenizer for chunk size estimation
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception:
            self.tokenizer = None

    def count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        if self.tokenizer:
            return len(self.tokenizer.encode(text))
        # Fallback: rough estimation for Turkish text
        return len(text) // 4

    def load_file(self, file_path: Path) -> LoadedDocument:
        """
        Load a single file.

        Args:
            file_path: Path to the file.

        Returns:
            LoadedDocument object.
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {suffix}")

        # Load content based on file type
        if suffix in {".md", ".txt"}:
            content = self._load_text(file_path)
        elif suffix == ".pdf":
            content = self._load_pdf(file_path)
        elif suffix == ".docx":
            content = self._load_docx(file_path)
        else:
            raise ValueError(f"No loader for {suffix}")

        # Create document
        doc_id = str(uuid.uuid4())
        doc = LoadedDocument(
            id=doc_id,
            filename=file_path.name,
            content=content,
            metadata={
                "source": str(file_path),
                "format": suffix,
                "filename": file_path.name
            }
        )

        # Chunk the document
        doc.chunks = self._chunk_document(doc)

        return doc

    def load_directory(
        self,
        directory: Path,
        recursive: bool = True
    ) -> List[LoadedDocument]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Path to the directory.
            recursive: Whether to search recursively.

        Returns:
            List of LoadedDocument objects.
        """
        directory = Path(directory)
        documents = []

        pattern = "**/*" if recursive else "*"

        for ext in self.SUPPORTED_EXTENSIONS:
            for file_path in directory.glob(f"{pattern}{ext}"):
                try:
                    doc = self.load_file(file_path)
                    documents.append(doc)
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")

        return documents

    def _load_text(self, file_path: Path) -> str:
        """Load plain text or markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF file."""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF is required for PDF support. Install with: pip install pymupdf")

        doc = fitz.open(file_path)
        text_parts = []

        for page in doc:
            text_parts.append(page.get_text())

        doc.close()
        return "\n\n".join(text_parts)

    def _load_docx(self, file_path: Path) -> str:
        """Load Word document."""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word support. Install with: pip install python-docx")

        doc = WordDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _chunk_document(self, doc: LoadedDocument) -> List[DocumentChunk]:
        """
        Split document into chunks.

        Args:
            doc: Document to chunk.

        Returns:
            List of DocumentChunk objects.
        """
        content = doc.content
        chunks = []

        # Split by paragraphs first
        paragraphs = content.split("\n\n")

        current_chunk = []
        current_tokens = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_tokens = self.count_tokens(para)

            # If single paragraph exceeds chunk size, split it
            if para_tokens > self.chunk_size:
                # Save current chunk if any
                if current_chunk:
                    chunk_text = "\n\n".join(current_chunk)
                    if self.count_tokens(chunk_text) >= self.min_chunk_size:
                        chunk_id = f"{doc.id}_chunk_{len(chunks)}"
                        chunks.append(DocumentChunk(
                            id=chunk_id,
                            content=chunk_text,
                            metadata={**doc.metadata, "chunk_index": len(chunks)}
                        ))
                    current_chunk = []
                    current_tokens = 0

                # Split large paragraph by sentences
                sentences = self._split_sentences(para)
                for sentence in sentences:
                    sentence_tokens = self.count_tokens(sentence)
                    if current_tokens + sentence_tokens > self.chunk_size:
                        if current_chunk:
                            chunk_text = " ".join(current_chunk)
                            if self.count_tokens(chunk_text) >= self.min_chunk_size:
                                chunk_id = f"{doc.id}_chunk_{len(chunks)}"
                                chunks.append(DocumentChunk(
                                    id=chunk_id,
                                    content=chunk_text,
                                    metadata={**doc.metadata, "chunk_index": len(chunks)}
                                ))
                        # Keep overlap
                        overlap_tokens = 0
                        overlap_idx = len(current_chunk)
                        while overlap_idx > 0 and overlap_tokens < self.chunk_overlap:
                            overlap_idx -= 1
                            overlap_tokens += self.count_tokens(current_chunk[overlap_idx])
                        current_chunk = current_chunk[overlap_idx:]
                        current_tokens = overlap_tokens

                    current_chunk.append(sentence)
                    current_tokens += sentence_tokens

            # Normal paragraph
            elif current_tokens + para_tokens > self.chunk_size:
                # Save current chunk
                if current_chunk:
                    chunk_text = "\n\n".join(current_chunk)
                    if self.count_tokens(chunk_text) >= self.min_chunk_size:
                        chunk_id = f"{doc.id}_chunk_{len(chunks)}"
                        chunks.append(DocumentChunk(
                            id=chunk_id,
                            content=chunk_text,
                            metadata={**doc.metadata, "chunk_index": len(chunks)}
                        ))

                # Keep some overlap from the end
                overlap_text = current_chunk[-1] if current_chunk else ""
                current_chunk = [overlap_text, para] if overlap_text else [para]
                current_tokens = self.count_tokens("\n\n".join(current_chunk))
            else:
                current_chunk.append(para)
                current_tokens += para_tokens

        # Don't forget the last chunk
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            if self.count_tokens(chunk_text) >= self.min_chunk_size:
                chunk_id = f"{doc.id}_chunk_{len(chunks)}"
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    content=chunk_text,
                    metadata={**doc.metadata, "chunk_index": len(chunks)}
                ))

        return chunks

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting for Turkish
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
